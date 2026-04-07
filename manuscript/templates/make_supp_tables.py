"""
make_supp_tables.py
Generate supplementary DOCX and CSV files for each chapter and place them
in the corresponding output/submission/<journal>/chNN/supp/ folder.

Supplementary content is defined here (mirroring QMD content) rather than
embedded in the manuscript body, consistent with the image workflow.

Usage:
    python templates/make_supp_tables.py            # all chapters
    python templates/make_supp_tables.py --chapter 1
    python templates/make_supp_tables.py --chapter 3
    python templates/make_supp_tables.py --chapter 4
"""
import sys
import csv
import json
import shutil
import argparse
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── helpers ────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def bold_row(row):
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
            if not para.runs:
                run = para.add_run(para.text)
                run.bold = True


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)


def save_table_docx(path: Path, caption: str, headers: list, rows: list):
    doc = Document()
    doc.add_paragraph()  # spacer

    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"

    # Header row
    hdr = tbl.rows[0]
    set_cell_bg(hdr.cells[0], "D9D9D9")
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, "D9D9D9")
        cell.text = h
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True

    # Data rows
    for r, row_data in enumerate(rows):
        tbl_row = tbl.rows[r + 1]
        for c, val in enumerate(row_data):
            tbl_row.cells[c].text = str(val)

    add_caption(doc, caption)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    print(f"  Saved: {path.name}")


# ── helpers for multi-section DOCX ─────────────────────────────────────────

def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(12)


def save_csv(path: Path, headers: list, rows: list):
    path.parent.mkdir(parents=True, exist_ok=True)
    with open(path, "w", newline="", encoding="utf-8") as f:
        w = csv.writer(f)
        w.writerow(headers)
        w.writerows(rows)
    print(f"  Saved: {path.name}")


# ── chapter supplementary file definitions ──────────────────────────────────

def _load_checklist(root: Path) -> list:
    """Load article_review_checklist.csv and return all rows as dicts."""
    path = root / "infrastructure_setup" / "manual_review" / "article_review_checklist.csv"
    with open(path, encoding="utf-8", errors="replace") as f:
        return list(csv.DictReader(f))


def _first_author(authors_str: str) -> str:
    if not authors_str:
        return ""
    return authors_str.split(",")[0].strip()


def _has_tag(tags_str: str, tag: str) -> str:
    return "Y" if tag in tags_str else "N"


def make_ch01_supp_files(supp_dir: Path):
    """CH_1 / CTS Supplementary Files S1–S5."""
    root = Path(__file__).parent.parent    # manuscript root (templates/../)

    # Load pipeline data
    all_rows = _load_checklist(root)
    # S3: decide-phase ML/XAI studies with NIH AI checklist applied
    ml_studies = [
        r for r in all_rows
        if r.get("ooda_phase", "").strip() == "decide"
        and r.get("nih_ai_score", "").strip()
    ]
    # S4: prediction model studies (decide + modeling crisp phase)
    probast_studies = [
        r for r in ml_studies
        if r.get("crisp_dm_phase", "").strip() == "modeling"
    ]

    # ── File S1: Full Boolean search strings ──────────────────────────────
    doc = Document()
    doc.add_paragraph()
    p = doc.add_heading("Supplementary File S1", level=1)
    doc.add_paragraph(
        "Full PubMed Boolean search strings for the Systematic Quantitative "
        "Literature Review (SQLR). Searches were conducted via PubMed API "
        "(January 2015–March 2026) across nine targeted strings organised by "
        "OODA-phase keyword ontology."
    )

    strings = [
        ("String 1 — Opioid ED Risk (core)",
         '("opioid use disorder"[MeSH] OR "opioid"[tiab] OR "OUD"[tiab]) AND '
         '("emergency department"[tiab] OR "ED visit"[tiab] OR "emergency room"[tiab]) AND '
         '("machine learning"[tiab] OR "predictive model"[tiab] OR "risk prediction"[tiab])'),
        ("String 2 — Polypharmacy & ADE",
         '("polypharmacy"[tiab] OR "adverse drug event"[MeSH] OR "drug-drug interaction"[tiab]) AND '
         '("machine learning"[tiab] OR "artificial intelligence"[tiab] OR "risk score"[tiab]) AND '
         '("claims"[tiab] OR "EHR"[tiab] OR "electronic health record"[tiab])'),
        ("String 3 — APCD / Claims Data",
         '("all-payer claims"[tiab] OR "all payer claims database"[tiab] OR "APCD"[tiab] OR '
         '"Medicare"[MeSH] OR "Medicaid"[MeSH]) AND '
         '("opioid"[tiab] OR "polypharmacy"[tiab]) AND '
         '("prediction"[tiab] OR "modeling"[tiab])'),
        ("String 4 — Pharmacogenomics (PGx)",
         '("pharmacogenomics"[MeSH] OR "pharmacogenetics"[tiab] OR "PGx"[tiab] OR '
         '"CYP2D6"[tiab] OR "OPRM1"[tiab] OR "gene-drug interaction"[tiab]) AND '
         '("opioid"[tiab] OR "polypharmacy"[tiab] OR "adverse drug"[tiab])'),
        ("String 5 — CPIC & Gene-Drug Guidelines",
         '("CPIC"[tiab] OR "Clinical Pharmacogenomics Implementation Consortium"[tiab] OR '
         '"metabolizer phenotype"[tiab] OR "poor metabolizer"[tiab]) AND '
         '("machine learning"[tiab] OR "clinical decision support"[tiab] OR "risk"[tiab])'),
        ("String 6 — XAI / Explainability",
         '("explainable artificial intelligence"[tiab] OR "explainable AI"[tiab] OR '
         '"SHAP"[tiab] OR "LIME"[tiab] OR "feature importance"[tiab] OR '
         '"interpretable machine learning"[tiab]) AND '
         '("opioid"[tiab] OR "polypharmacy"[tiab] OR "pharmacogenomic"[tiab] OR "clinical"[tiab])'),
        ("String 7 — Gradient Boosting Methods",
         '("gradient boosting"[tiab] OR "XGBoost"[tiab] OR "CatBoost"[tiab] OR '
         '"random forest"[tiab] OR "LightGBM"[tiab]) AND '
         '("opioid"[tiab] OR "polypharmacy"[tiab] OR "adverse drug event"[tiab] OR '
         '"pharmacogenomic"[tiab])'),
        ("String 8 — Causal / Temporal Modeling",
         '("causal inference"[tiab] OR "target leakage"[tiab] OR "temporal validation"[tiab] OR '
         '"dynamic time warping"[tiab] OR "DTW"[tiab] OR "process mining"[tiab]) AND '
         '("opioid"[tiab] OR "polypharmacy"[tiab] OR "clinical prediction"[tiab])'),
        ("String 9 — Clinical Decision Support Deployment",
         '("clinical decision support"[MeSH] OR "CDS"[tiab] OR "point-of-care"[tiab] OR '
         '"serverless"[tiab] OR "cloud"[tiab]) AND '
         '("pharmacogenomic"[tiab] OR "opioid risk"[tiab] OR "polypharmacy risk"[tiab]) AND '
         '("machine learning"[tiab] OR "artificial intelligence"[tiab])'),
    ]

    for title, string in strings:
        add_heading(doc, title, level=2)
        p = doc.add_paragraph(string)
        p.runs[0].font.name = "Courier New"
        p.runs[0].font.size = Pt(9)

    path = supp_dir / "File_S1.docx"
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    print(f"  Saved: {path.name}")

    # ── File S2: PRISMA 2020 Checklist ────────────────────────────────────
    doc = Document()
    doc.add_paragraph()
    doc.add_heading("Supplementary File S2", level=1)
    doc.add_paragraph(
        "PRISMA 2020 Checklist for Systematic Quantitative Literature Review: "
        "Bridging Explainable Artificial Intelligence and Pharmacogenomics for "
        "Opioid and Polypharmacy Risk Prediction."
    )
    save_table_docx(
        path=supp_dir / "File_S2.docx",
        caption="Table S2. PRISMA 2020 Checklist.",
        headers=["Section", "Item #", "Checklist Item", "Location in Manuscript"],
        rows=[
            ["TITLE",        "1",  "Identify the report as a systematic review.",                                                  "Title"],
            ["ABSTRACT",     "2",  "See the PRISMA 2020 for Abstracts checklist.",                                                 "Abstract"],
            ["INTRODUCTION", "3",  "Describe the rationale for the review.",                                                       "§1.1"],
            ["INTRODUCTION", "4",  "Provide an explicit statement of the objectives.",                                             "§1.4"],
            ["METHODS",      "5",  "Specify the inclusion and exclusion criteria.",                                                "§2.2"],
            ["METHODS",      "6",  "Specify all databases, registers, websites, and other sources searched.",                      "§2.3"],
            ["METHODS",      "7",  "Present the full search strategies for all databases.",                                        "File S1"],
            ["METHODS",      "8",  "Specify the methods used to decide whether a study met the inclusion criteria.",               "§2.2"],
            ["METHODS",      "9",  "Specify the methods used to collect data from reports.",                                       "§2.4"],
            ["METHODS",      "10", "List and define all outcomes for which data were sought.",                                     "§2.2"],
            ["METHODS",      "11", "Describe the methods used to assess risk of bias.",                                            "§2.5"],
            ["METHODS",      "12", "Specify all measures used to synthesize results.",                                             "§2.6"],
            ["METHODS",      "13", "Describe any methods to identify selective reporting.",                                        "§2.5"],
            ["METHODS",      "14", "Describe any sensitivity analyses conducted.",                                                  "§2.6"],
            ["METHODS",      "15", "Describe any methods to assess certainty of evidence.",                                        "§2.5"],
            ["RESULTS",      "16", "Describe the results of the search and selection process (PRISMA flow diagram).",              "§3.1 / Fig. 1"],
            ["RESULTS",      "17", "Cite each included study and present its characteristics.",                                    "§3.2 / File S3"],
            ["RESULTS",      "18", "Present assessments of risk of bias for each included study.",                                 "§3.3 / File S4"],
            ["RESULTS",      "19", "For all outcomes, present a summary of available data.",                                       "§3.2–3.4"],
            ["RESULTS",      "20", "Present results of all statistical syntheses.",                                                "§3.4"],
            ["RESULTS",      "21", "Present results of any investigations of causes of heterogeneity.",                           "§3.4"],
            ["DISCUSSION",   "22", "Provide a general interpretation of the results.",                                             "§4"],
            ["DISCUSSION",   "23", "Discuss any limitations of the evidence.",                                                     "§4.3"],
            ["DISCUSSION",   "24", "Describe the implications of the results.",                                                    "§4.4"],
            ["OTHER",        "25", "Provide registration information for the review.",                                             "Abstract (PROSPERO)"],
            ["OTHER",        "26", "Indicate where the review protocol can be accessed.",                                          "Not pre-registered"],
            ["OTHER",        "27", "Describe sources of financial or non-financial support.",                                      "Funding"],
        ],
    )

    # ── File S3: Data Extraction CSV (151 decide-phase ML/XAI studies) ──────
    s3_headers = [
        "Study_ID", "First_Author", "Year", "Title", "DOI",
        "OODA_Phase", "CRISP_DM_Phase", "OODA_CRISP_Label",
        "NIH_AI_Score",
        "Has_Performance_Metrics", "Has_External_Validation",
        "Has_Explainability", "Has_Model_Transparency",
        "Has_Bias_Fairness", "Has_Clinical_Utility",
        "Has_Safety_Monitoring", "Has_Regulatory_Ethics",
        "OP_Perf_Tags", "Has_PDF", "Notes",
    ]
    s3_rows = []
    for r in ml_studies:
        tags = r.get("nih_ai_tags", "")
        s3_rows.append([
            r.get("article_id", "").strip(),
            _first_author(r.get("authors", "")),
            r.get("pub_year", "").strip(),
            r.get("title", "").strip(),
            r.get("doi", "").strip(),
            r.get("ooda_phase", "").strip(),
            r.get("crisp_dm_phase", "").strip(),
            r.get("ooda_crisp_label", "").strip(),
            r.get("nih_ai_score", "").strip(),
            _has_tag(tags, "performance_metrics"),
            _has_tag(tags, "external_validation"),
            _has_tag(tags, "explainability"),
            _has_tag(tags, "model_transparency"),
            _has_tag(tags, "bias_fairness"),
            _has_tag(tags, "clinical_utility"),
            _has_tag(tags, "safety_monitoring"),
            _has_tag(tags, "regulatory_ethics"),
            r.get("op_perf_tags", "").strip(),
            r.get("has_pdf", "").strip(),
            r.get("notes", "").strip(),
        ])
    save_csv(path=supp_dir / "File_S3.csv", headers=s3_headers, rows=s3_rows)
    print(f"  ({len(s3_rows)} studies)")

    # ── File S4: PROBAST Assessments DOCX (56 decide+modeling studies) ────
    def _probast_d1(tags):  # Participants: bias_fairness → L else U
        return "L" if "bias_fairness" in tags else "U"

    def _probast_d2(tags):  # Predictors: data_reporting → L else U
        return "L" if "data_reporting" in tags else "U"

    def _probast_d3(tags):  # Outcome: performance_metrics → L else H
        return "L" if "performance_metrics" in tags else "H"

    def _probast_d4(tags):  # Analysis: transparency or explainability → L else H
        return "L" if ("model_transparency" in tags or "explainability" in tags) else "H"

    def _probast_overall(d1, d2, d3, d4):
        return "H" if "H" in (d1, d2, d3, d4) else ("U" if "U" in (d1, d2, d3, d4) else "L")

    def _concern(tags):
        return "Low" if ("clinical_utility" in tags and "safety_monitoring" in tags) else "Moderate"

    s4_rows = []
    for r in probast_studies:
        tags = r.get("nih_ai_tags", "")
        d1, d2, d3, d4 = _probast_d1(tags), _probast_d2(tags), _probast_d3(tags), _probast_d4(tags)
        s4_rows.append([
            r.get("article_id", "").strip(),
            f"{_first_author(r.get('authors', ''))} ({r.get('pub_year','').strip()})",
            r.get("title", "").strip()[:100],
            d1, d2, d3, d4,
            _probast_overall(d1, d2, d3, d4),
            _concern(tags),
        ])
    save_table_docx(
        path=supp_dir / "File_S4.docx",
        caption=(
            f"Table S4. PROBAST risk of bias and applicability assessments for "
            f"{len(s4_rows)} prediction model studies (decide+modeling phase). "
            "D1 Participants, D2 Predictors, D3 Outcome, D4 Analysis. "
            "L = Low, H = High, U = Unclear. "
            "Domain ratings derived from NIH AI checklist tags."
        ),
        headers=[
            "Study_ID", "First Author (Year)", "Title (truncated)",
            "D1", "D2", "D3", "D4", "Overall ROB", "Overall Concern",
        ],
        rows=s4_rows,
    )
    print(f"  ({len(s4_rows)} studies)")

    # ── File S5: Evidence Map CSV (all 5,839 included studies) ────────────
    s5_headers = [
        "Rank", "Study_ID", "First_Author", "Year", "Title", "DOI",
        "OODA_Phase", "CRISP_DM_Phase", "OODA_CRISP_Label",
        "Composite_Score", "PyTextRank_Score", "Combined_Score",
        "NIH_AI_Score", "NIH_AI_Tags", "OP_Perf_Tags",
        "Has_PDF", "Notes",
    ]
    s5_rows = []
    for r in all_rows:
        s5_rows.append([
            r.get("rank", "").strip(),
            r.get("article_id", "").strip(),
            _first_author(r.get("authors", "")),
            r.get("pub_year", "").strip(),
            r.get("title", "").strip(),
            r.get("doi", "").strip(),
            r.get("ooda_phase", "").strip(),
            r.get("crisp_dm_phase", "").strip(),
            r.get("ooda_crisp_label", "").strip(),
            r.get("composite_score", "").strip(),
            r.get("pytextrank_score", "").strip(),
            r.get("combined_score", "").strip(),
            r.get("nih_ai_score", "").strip(),
            r.get("nih_ai_tags", "").strip(),
            r.get("op_perf_tags", "").strip(),
            r.get("has_pdf", "").strip(),
            r.get("notes", "").strip(),
        ])
    save_csv(path=supp_dir / "File_S5.csv", headers=s5_headers, rows=s5_rows)
    print(f"  ({len(s5_rows)} studies)")


# ── chapter table definitions ───────────────────────────────────────────────

def make_ch03_supp_figures(supp_dir: Path):
    """CH_3 / CTS-2026-0196 supplementary figures (PNG, embedded for CTS)."""
    root = Path(__file__).parent.parent
    fig_dir = root / "figures" / "ch03"

    figures = [
        ("fig_shap_pdp.png",             "Figure_S1.png"),
        ("fig_dtw_pathways.png",         "Figure_S2.png"),
        ("fig_trajectories_heatmap.png", "Figure_S3.png"),
    ]

    supp_dir.mkdir(parents=True, exist_ok=True)
    for src_name, dest_name in figures:
        src = fig_dir / src_name
        dest = supp_dir / dest_name
        if src.exists():
            shutil.copy2(src, dest)
            print(f"  Copied: {src_name} -> {dest_name}")
        else:
            print(f"  WARNING: source not found: {src_name}")


def make_ch04_tables(supp_dir: Path):
    """CH_4 / PSP-2026-0109 supplementary tables."""

    # Table S2 — Top deprescribing targets by Intervention Rate
    save_table_docx(
        path=supp_dir / "Table_S2.docx",
        caption=(
            "Table S2. Top deprescribing targets by Intervention Rate (IR) with "
            "clinical code mapping. NDC prefixes are representative; full NDC varies "
            "by manufacturer. IR values are per-patient probability shifts normalized "
            "by case prevalence."
        ),
        headers=["Drug", "NDC (prefix)", "Top ICD-10", "STOPP/Beers", "IR (65\u201374)"],
        rows=[
            ["Simvastatin",  "00006-0726", "E78.5",       "Beers (CYP3A4)",   "7.0 \u00d7 10\u207b\u2074"],
            ["Furosemide",   "00469-2620", "I50.9, N18.x","STOPP D4",         "2.0 \u00d7 10\u207b\u2074"],
            ["Alprazolam",   "00009-0029", "F41.1, G40.x","Beers CNS; STOPP D6","1.0 \u00d7 10\u207b\u2074"],
            ["Levofloxacin", "00069-7520", "J18.9",       "\u2014",           "Top-5 (all bands)"],
            ["Lorazepam",    "00069-0081", "F41.1, G40.x","Beers CNS",        "Top-5 (all bands)"],
        ],
    )

    # Table S1 — 115 synergistic drug pairs with IE and 95% CI
    root = Path(__file__).parent.parent
    pairs_path = root / "data" / "ffa_synergy_pairs.json"
    if not pairs_path.exists():
        print(f"  WARNING: {pairs_path} not found — run extract_ffa_table_s1.py first")
    else:
        pairs = json.loads(pairs_path.read_text(encoding="utf-8"))
        s1_headers = [
            "Drug A", "Drug B", "Age Band", "IE Score",
            "IR\u2090 (\u00d710\u207b\u2074)", "IR\u1d07 (\u00d710\u207b\u2074)",
            "Support\u207a", "95% CI"
        ]
        s1_rows = []
        for p in pairs:
            ci = (f"{p['ci_lo']:.2f}\u2013{p['ci_hi']:.2f}"
                  if p.get("ci_lo") is not None else "\u2014")
            ir_a = f"{p['ir_a']*1e4:.2f}" if p['ir_a'] else "0.00"
            ir_b = f"{p['ir_b']*1e4:.2f}" if p['ir_b'] else "0.00"
            s1_rows.append([
                p["drug_a"], p["drug_b"], p["age_band"],
                f"{p['ie']:.3f}", ir_a, ir_b,
                str(p["support_pos"]), ci,
            ])
        save_table_docx(
            path=supp_dir / "Table_S1.docx",
            caption=(
                f"Table S1. All {len(pairs)} synergistic drug-drug interaction pairs "
                "identified in non-opioid-related ED cohort (non_opioid_ed) across "
                "geriatric age bands (65\u201374, 75\u201384, 85\u2013114). "
                "IE = Interaction Effect (AXP inner lift, IE > 1.0 = synergistic). "
                "IR = Causal Responsibility (Intervention Rate, scaled \u00d710\u207b\u2074). "
                "Support\u207a = count of class-1 AXP rules containing both drugs. "
                "95% CI shown for top 5 pairs (bootstrap, n=1,000 resamples, non_opioid_ed/85\u2013114)."
            ),
            headers=s1_headers,
            rows=s1_rows,
        )
        save_csv(
            path=supp_dir / "Table_S1.csv",
            headers=["Drug_A","Drug_B","Age_Band","IE","IR_A_1e4","IR_B_1e4","Support_Pos","CI_95"],
            rows=[[p["drug_a"],p["drug_b"],p["age_band"],p["ie"],
                   p["ir_a"]*1e4,p["ir_b"]*1e4,p["support_pos"],
                   (f"{p['ci_lo']:.2f}\u2013{p['ci_hi']:.2f}" if p.get("ci_lo") else "")
                   ] for p in pairs],
        )
        print(f"  Table S1: {len(pairs)} pairs -> Table_S1.docx + Table_S1.csv")


# ── entry point ─────────────────────────────────────────────────────────────

CHAPTER_MAP = {
    1: ("cts",     "ch01", make_ch01_supp_files),
    3: ("cts",     "ch03", make_ch03_supp_figures),
    4: ("cpt_psp", "ch04", make_ch04_tables),
}


def main():
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--chapter", type=int, default=0,
                        help="Chapter number (0 = all chapters with supp tables)")
    args = parser.parse_args()

    root = Path(__file__).parent.parent
    chapters = ([args.chapter] if args.chapter else list(CHAPTER_MAP.keys()))

    for ch in chapters:
        if ch not in CHAPTER_MAP:
            print(f"  Chapter {ch}: no supplementary tables defined — skipped")
            continue
        journal, ch_dir, builder = CHAPTER_MAP[ch]
        supp_dir = root / "output" / "submission" / journal / ch_dir / "supp"
        print(f"\n==> CH_{ch} supplementary tables -> {supp_dir}")
        builder(supp_dir)


if __name__ == "__main__":
    main()
