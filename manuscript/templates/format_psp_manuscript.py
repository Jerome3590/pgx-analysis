"""
format_psp_manuscript.py
PSP-specific DOCX post-processor (replaces Pass 2 insert_docx_images.py).

For CPT:PSP submission, figures must be uploaded as SEPARATE files — NOT
embedded in the manuscript. This script:

  1. Replaces every [IMAGE: path] placeholder paragraph with a callout
     paragraph styled as:  [Figure N near here]
  2. Collects every [LEGEND:] paragraph (emitted by suppress_images_psp.lua),
     strips the [LEGEND:] prefix, and moves them to a "Figure Legends" section
     appended at the end of the document (after References), per PSP style guide.

Usage:
    python templates/format_psp_manuscript.py <file.docx> --chapter N
"""
import re
import sys
import argparse
import copy
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt

# ── Placeholder patterns ────────────────────────────────────────────────────

IMAGE_RE  = re.compile(r'^\[IMAGE:(.+?)(?::[\d.]+%?)?\]$')
LEGEND_RE = re.compile(r'^\[LEGEND:\]\s*')

# ── Figure filename → number maps (must match export_figures_psp.py order) ─

SUPP_FIGURE_MAP = {
    2: {"pgx_architecture_analysis.png": 1},
    4: {"fig_shap_pdp.png": 1, "fig_trajectories.png": 2},
    5: {},
}

FIGURE_MAP = {
    2: {
        "pgx_architecture_clinical_ooda_loop.png": 1,
        "pgx_architecture_pipeline.png":           2,
        "pgx_architecture_consensus_filter.png":   3,
        "pgx_architecture_risk_dashboard.png":     4,
        "fig_attrition.png":                       5,
    },
    4: {
        "fig_shap.png":    1,
        "fig_network.png": 2,
        "fig_ir.png":      3,
        "fig_zcode.png":   4,
    },
    5: {
        "pgx_architecture_risk_dashboard.png": 1,
        "fig_imputation.png":                  2,
        "pgx_dashboard.png":                   3,
        "fig_latency.png":                     4,
    },
}


def get_text(para) -> str:
    return "".join(r.text for r in para.runs).strip()


def para_starts_with_legend(para) -> bool:
    text = get_text(para)
    return text.startswith("[LEGEND:]")


def strip_legend_marker(para) -> None:
    """Remove the leading [LEGEND:] text from the first run of a paragraph."""
    for run in para.runs:
        if "[LEGEND:]" in run.text:
            run.text = run.text.replace("[LEGEND:]", "").lstrip()
            return


def resolve_figure_num(path_str: str, ch_map: dict) -> int | None:
    filename = Path(path_str.strip()).name
    return ch_map.get(filename)


def add_heading(doc: Document, text: str) -> None:
    """Append a bold heading paragraph at the end of the document body."""
    body = doc.element.body
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    bold = OxmlElement("w:b")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "24")  # 12pt
    rpr.append(bold)
    rpr.append(sz)
    r.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    p.append(r)
    # Insert before the last sectPr if present
    last = list(body)[-1]
    if last.tag == qn("w:sectPr"):
        last.addprevious(p)
    else:
        body.append(p)


def replace_runs_text(para_elem, new_text: str, italic: bool = False) -> None:
    """Clear all runs in a paragraph and set the first run to new_text."""
    from docx.text.paragraph import Paragraph
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    runs = para_elem.findall(f"{{{ns}}}r")
    # Clear all runs
    for r in runs:
        for t in r.findall(f"{{{ns}}}t"):
            t.text = ""
    if runs:
        # Set text on first run; set italic if requested
        t_elem = runs[0].find(f"{{{ns}}}t")
        if t_elem is None:
            t_elem = OxmlElement("w:t")
            runs[0].append(t_elem)
        t_elem.text = new_text
        t_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        if italic:
            rpr = runs[0].find(f"{{{ns}}}rPr")
            if rpr is None:
                rpr = OxmlElement("w:rPr")
                runs[0].insert(0, rpr)
            i_elem = OxmlElement("w:i")
            rpr.append(i_elem)
    else:
        # No existing runs — create one
        r = OxmlElement("w:r")
        t_elem = OxmlElement("w:t")
        t_elem.text = new_text
        t_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        if italic:
            rpr = OxmlElement("w:rPr")
            i_elem = OxmlElement("w:i")
            rpr.append(i_elem)
            r.append(rpr)
        r.append(t_elem)
        para_elem.append(r)


def para_elem_text(para_elem) -> str:
    """Get full text of a w:p element by concatenating all w:t descendants."""
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    return "".join(t.text or "" for t in para_elem.iter(f"{{{ns}}}t")).strip()


def is_heading1(elem) -> bool:
    """True if the w:p element uses Heading 1 style."""
    if elem.tag != qn("w:p"):
        return False
    pPr = elem.find(qn("w:pPr"))
    if pPr is None:
        return False
    pStyle = pPr.find(qn("w:pStyle"))
    if pStyle is None:
        return False
    val = pStyle.get(qn("w:val"), "").lower()
    return "heading1" in val or val == "1"


def move_study_highlights(doc: Document) -> None:
    """Move Study Highlights section to just before the Introduction heading.

    PSP/CPT requires Study Highlights to appear after the abstract and before
    the main text body.  In the QMD source it sits near the end of the document
    for PDF rendering; here we relocate it for the DOCX submission package.
    """
    body     = doc.element.body
    children = list(body)

    def elem_text(elem) -> str:
        return "".join(t.text or "" for t in elem.iter(qn("w:t"))).strip()

    # 1. Find the Study Highlights Heading 1
    sh_start = None
    for i, elem in enumerate(children):
        if is_heading1(elem) and "Study Highlights" in elem_text(elem):
            sh_start = i
            break
    if sh_start is None:
        print("  Study Highlights heading not found — skipped")
        return

    # 2. Collect paragraphs through the next Heading 1 (exclusive)
    sh_end = len(children)
    for i in range(sh_start + 1, len(children)):
        if is_heading1(children[i]):
            sh_end = i
            break

    # 3. Find the Introduction Heading 1 (must appear BEFORE Study Highlights)
    intro_elem = None
    for i in range(sh_start):
        if is_heading1(children[i]) and "Introduction" in elem_text(children[i]):
            intro_elem = children[i]
            break
    if intro_elem is None:
        print("  Introduction heading not found before Study Highlights — skipped")
        return

    # 4. Deep-copy the Study Highlights block, then remove from current location
    sh_elems = [copy.deepcopy(children[i]) for i in range(sh_start, sh_end)]
    for i in range(sh_end - 1, sh_start - 1, -1):
        body.remove(children[i])

    # 5. Re-fetch Introduction element (index may have shifted after removal)
    intro_elem = None
    for elem in body:
        if is_heading1(elem) and "Introduction" in elem_text(elem):
            intro_elem = elem
            break
    if intro_elem is None:
        print("  Introduction heading lost after Study Highlights removal — skipped")
        return

    # 6. Insert Study Highlights block just before Introduction.
    #    Forward iteration with addprevious(intro_elem) preserves source order:
    #    H1 heading lands first, paragraphs follow in document order.
    for elem in sh_elems:
        intro_elem.addprevious(elem)
    print(f"  Study Highlights ({sh_end - sh_start} paragraphs) moved before Introduction")


def format_psp(docx_path: Path, chapter: int) -> None:
    ch_map = FIGURE_MAP.get(chapter, {})
    if not ch_map:
        print(f"  WARNING: no figure map for chapter {chapter} — placeholders left as-is")

    doc = Document(str(docx_path))
    body = doc.element.body
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    # Iterate ALL w:p elements recursively (pandoc wraps figures in w:tbl/w:tr/w:tc)
    legend_elems        = []  # deep copies of caption paragraphs
    legend_to_remove    = []  # (parent, elem) to remove from table cells
    replaced = 0
    collected = 0

    for para_elem in list(body.iter(f"{{{ns}}}p")):
        text = para_elem_text(para_elem)
        m = IMAGE_RE.match(text)
        if not m:
            continue

        fig_num = resolve_figure_num(m.group(1), ch_map)
        if fig_num:
            label = f"Figure {fig_num}"
        else:
            supp_map = SUPP_FIGURE_MAP.get(chapter, {})
            supp_num = resolve_figure_num(m.group(1), supp_map)
            label = f"Figure S{supp_num}" if supp_num else Path(m.group(1)).stem
        replace_runs_text(para_elem, f"[{label} near here]", italic=True)
        replaced += 1
        print(f"  [{label} near here]  ← {Path(m.group(1)).name}")

        # Caption is the next w:p sibling inside the same w:tc parent
        parent = para_elem.getparent()
        if parent is not None and parent.tag == qn("w:tc"):
            siblings = list(parent)
            idx = siblings.index(para_elem)
            for sibling in siblings[idx + 1:]:
                if sibling.tag == f"{{{ns}}}p":
                    legend_elems.append(copy.deepcopy(sibling))
                    legend_to_remove.append((parent, sibling))
                    collected += 1
                    break

    # Remove caption paragraphs from their table cells
    for parent, elem in legend_to_remove:
        parent.remove(elem)

    # Append Figure Legends section at the end of the body
    if legend_elems:
        add_heading(doc, "Figure Legends")
        last = list(body)[-1]
        sect = last if last.tag == qn("w:sectPr") else None
        for legend_elem in legend_elems:
            if sect is not None:
                sect.addprevious(legend_elem)
            else:
                body.append(legend_elem)

    # Move Study Highlights section to before Introduction (PSP/CPT submission order)
    move_study_highlights(doc)

    doc.save(str(docx_path))
    print(f"  {docx_path.name}: {replaced} callout(s) inserted, "
          f"{collected} legend(s) moved to Figure Legends section")


def main():
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("docx", help="Path to the .docx file")
    parser.add_argument("--chapter", type=int, required=True,
                        help="Chapter number (for figure→number mapping)")
    args = parser.parse_args()

    docx_path = Path(args.docx).resolve()
    if not docx_path.exists():
        print(f"ERROR: {docx_path} not found")
        sys.exit(1)

    print(f"Processing (PSP): {docx_path.name}  [chapter {args.chapter}]")
    format_psp(docx_path, args.chapter)


if __name__ == "__main__":
    main()
