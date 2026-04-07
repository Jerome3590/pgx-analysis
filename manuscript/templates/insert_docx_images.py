"""
insert_docx_images.py
Second pass of the two-pass DOCX build:
  - Finds [IMAGE:path] placeholder paragraphs left by suppress_images.lua
  - Loads the source image file, computes scaled dimensions
  - Inserts the image inline (wp:inline) at the placeholder location
  - Preserves paragraph style and surrounding content

Usage:
    python templates/insert_docx_images.py <file.docx> [--chapter-dir CH_1]
    python templates/insert_docx_images.py output/edits/cts/ch01_cts_draft.docx --chapter-dir CH_1

The chapter-dir is the directory containing the .qmd source (used to resolve
relative image paths like ../figures/ch01/fig.png).
If omitted, paths are resolved relative to the current working directory.
"""
import sys
import re
import argparse
import subprocess
from pathlib import Path

from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image as PILImage

WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"

MAX_W_IN   = 6.0   # hard cap: DOCX text area width
MAX_H_IN   = 8.5   # hard cap: Letter page minus margins
DOCX_TW_IN = 6.3   # actual DOCX text width (8.5" - 1" - 1" - 0.1" gutter)

PLACEHOLDER = re.compile(r'^\[IMAGE:(.+?)(?::([\d.]+%?))?\]$')


def parse_width(width_str: str | None) -> float | None:
    """Convert a QMD width attribute ('85%', '100%', '4in') to inches, or None."""
    if not width_str:
        return None
    width_str = width_str.strip()
    if width_str.endswith('%'):
        pct = float(width_str[:-1]) / 100.0
        return min(pct * DOCX_TW_IN, MAX_W_IN)
    if width_str.endswith('in'):
        return min(float(width_str[:-2]), MAX_W_IN)
    return None


def scaled_dims(img_path: Path, target_w: float | None = None) -> tuple[float, float]:
    """Return (width_in, height_in) at target_w if given, else scaled to MAX_W."""
    with PILImage.open(img_path) as im:
        w_px, h_px = im.size
        dpi_info = im.info.get('dpi', (96, 96))
        dpi = dpi_info[0] if isinstance(dpi_info, tuple) else 96
        if dpi < 10:
            dpi = 96
    natural_w = w_px / dpi
    natural_h = h_px / dpi
    aspect = natural_h / natural_w if natural_w else 1.0

    # Use target width from QMD attribute, falling back to natural/max
    w_in = target_w if target_w else min(natural_w, MAX_W_IN)
    h_in = w_in * aspect

    # Safety cap: never overflow the page
    if h_in > MAX_H_IN:
        ratio = MAX_H_IN / h_in
        w_in, h_in = w_in * ratio, MAX_H_IN
    return w_in, h_in


def inline_to_anchor(run) -> None:
    """Convert the just-inserted wp:inline image to wp:anchor with Top and Bottom wrapping."""
    drawing = run._element.find(qn('w:drawing'))
    if drawing is None:
        return
    inline = drawing.find(f'{{{WP_NS}}}inline')
    if inline is None:
        return

    # Rename tag inline → anchor
    inline.tag = f'{{{WP_NS}}}anchor'

    # Set anchor-required attributes
    for attr in ('distT', 'distB', 'distL', 'distR'):
        inline.attrib.pop(attr, None)
    inline.set('distT', '0')
    inline.set('distB', '0')
    inline.set('distL', '114300')
    inline.set('distR', '114300')
    inline.set('simplePos', '0')
    inline.set('relativeHeight', '251658240')
    inline.set('behindDoc', '0')
    inline.set('locked', '0')
    inline.set('layoutInCell', '1')
    inline.set('allowOverlap', '1')

    # Build positional children to prepend before wp:extent
    sp = OxmlElement('wp:simplePos')
    sp.set('x', '0'); sp.set('y', '0')

    ph = OxmlElement('wp:positionH')
    ph.set('relativeFrom', 'column')
    ah = OxmlElement('wp:align')
    ah.text = 'center'
    ph.append(ah)

    pv = OxmlElement('wp:positionV')
    pv.set('relativeFrom', 'paragraph')
    po = OxmlElement('wp:posOffset')
    po.text = '0'
    pv.append(po)

    # Insert simplePos, positionH, positionV at front (before extent)
    inline.insert(0, pv)
    inline.insert(0, ph)
    inline.insert(0, sp)

    # Insert wrapTopAndBottom just before wp:docPr
    doc_pr = inline.find(f'{{{WP_NS}}}docPr')
    wrap = OxmlElement('wp:wrapTopAndBottom')
    if doc_pr is not None:
        inline.insert(list(inline).index(doc_pr), wrap)
    else:
        inline.append(wrap)


def clear_paragraph(para) -> None:
    """Remove all runs and text from a paragraph, keeping its style."""
    for child in list(para._p):
        if child.tag in (qn('w:r'), qn('w:hyperlink'), qn('w:ins'), qn('w:del')):
            para._p.remove(child)


def all_paragraphs(doc):
    """Yield every paragraph in the document, including those inside table cells."""
    yield from doc.paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def pdf_to_png(pdf_path: Path) -> Path | None:
    """Convert a PDF file to PNG using ImageMagick. Returns the PNG path or None."""
    png_path = pdf_path.with_suffix('.png')
    if png_path.exists():
        return png_path
    try:
        result = subprocess.run(
            ['magick', '-density', '300', f'{pdf_path}[0]', '-quality', '95', str(png_path)],
            capture_output=True, timeout=30
        )
        if result.returncode == 0 and png_path.exists():
            return png_path
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass
    return None


def resolve_image(src: str, docx_dir: Path, chapter_dir: Path) -> Path | None:
    """Locate image file; for .png paths, also try PDF→PNG conversion as fallback."""
    search_bases = (docx_dir, chapter_dir, docx_dir.parent, docx_dir.parent.parent)
    for base in search_bases:
        candidate = (base / src).resolve()
        if candidate.exists():
            return candidate
    # PNG not found: look for a PDF sibling in figure-pdf/ directories
    if src.endswith('.png'):
        src_pdf = src.replace('figure-docx', 'figure-pdf').replace('.png', '.pdf')
        for base in search_bases:
            candidate_pdf = (base / src_pdf).resolve()
            if candidate_pdf.exists():
                png = pdf_to_png(candidate_pdf)
                if png:
                    return png
    return None


def insert_images(docx_path: Path, chapter_dir: Path) -> int:
    doc = Document(str(docx_path))
    docx_dir = docx_path.parent
    inserted = 0

    for para in all_paragraphs(doc):
        text = para.text.strip()
        m = PLACEHOLDER.match(text)
        if not m:
            continue

        src = m.group(1).strip()
        width_attr = m.group(2)  # may be None
        target_w = parse_width(width_attr)
        img_path = resolve_image(src, docx_dir, chapter_dir)
        if img_path is None:
            print(f"  WARNING: image not found: {src}")
            continue

        try:
            w_in, h_in = scaled_dims(img_path, target_w)
        except Exception as e:
            print(f"  WARNING: could not read {img_path}: {e}")
            continue

        clear_paragraph(para)
        run = para.add_run()
        run.add_picture(str(img_path), width=Inches(w_in), height=Inches(h_in))
        inline_to_anchor(run)
        inserted += 1
        print(f"  Inserted {img_path.name} at {w_in:.2f}\" × {h_in:.2f}\" [anchor]")

    doc.save(str(docx_path))
    return inserted


def main():
    parser = argparse.ArgumentParser(description="Insert images into a DOCX after Quarto render")
    parser.add_argument("docx", help="Path to the .docx file")
    parser.add_argument("--chapter-dir", default=None,
                        help="Chapter source directory (for resolving relative image paths)")
    args = parser.parse_args()

    docx_path = Path(args.docx).resolve()
    if not docx_path.exists():
        print(f"ERROR: {docx_path} not found")
        sys.exit(1)

    if args.chapter_dir:
        chapter_dir = Path(args.chapter_dir).resolve()
    else:
        chapter_dir = Path.cwd()

    print(f"Processing: {docx_path.name}")
    print(f"Image base: {chapter_dir}")
    n = insert_images(docx_path, chapter_dir)
    print(f"  Done: {n} image(s) inserted")


if __name__ == "__main__":
    main()
