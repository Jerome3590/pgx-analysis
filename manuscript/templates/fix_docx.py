"""
Post-process Quarto-generated DOCX files:
  1. Left-justify all table cell text (override Pandoc center/justify attributes)
  2. Mark all field codes dirty so Word recalculates TOC page numbers on first open

Note: image insertion and sizing is handled upstream by insert_docx_images.py,
which produces wp:anchor floats with Top-and-Bottom text wrapping.

Usage:
    python templates/fix_docx.py output/edits/cts/ch01_cts_draft.docx
    python templates/fix_docx.py output/edits/    # process all *.docx in tree
"""
import sys
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def mark_fields_dirty(doc: Document) -> int:
    """Mark all field-begin chars as dirty so Word recalculates TOC/fields on open."""
    count = 0
    for elem in doc.element.body.iter(qn('w:fldChar')):
        if elem.get(qn('w:fldCharType')) == 'begin':
            elem.set(qn('w:dirty'), 'true')
            count += 1
    return count


def fix_table_alignment(doc: Document) -> int:
    fixed = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    pPr = para._p.get_or_add_pPr()
                    jc = pPr.find(qn("w:jc"))
                    if jc is None:
                        jc = OxmlElement("w:jc")
                        pPr.append(jc)
                    jc.set(qn("w:val"), "left")
                    fixed += 1
    return fixed


def process_file(path: Path) -> None:
    doc = Document(str(path))
    cells  = fix_table_alignment(doc)
    fields = mark_fields_dirty(doc)
    doc.save(str(path))
    print(f"  {path.name}: {cells} table cells left-justified, {fields} fields marked dirty")


def main():
    if len(sys.argv) < 2:
        print("Usage: fix_docx.py <file.docx|directory>")
        sys.exit(1)

    target = Path(sys.argv[1])
    if target.is_file():
        process_file(target)
    elif target.is_dir():
        files = list(target.rglob("*.docx"))
        if not files:
            print(f"No .docx files found in {target}")
            sys.exit(1)
        for f in files:
            process_file(f)
    else:
        print(f"Not found: {target}")
        sys.exit(1)


if __name__ == "__main__":
    main()
